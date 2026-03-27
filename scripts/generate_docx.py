#!/usr/bin/env python3
# Copyright 2026 Adam Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generate DOCX files from all Markdown sources in the Rare Disease
Diagnostic Agent docs/ directory.

Applies NVIDIA / HCLS AI Factory branding (VCP palette).  Handles
headings, bold/italic runs, bullet and numbered lists, tables, code
blocks, and horizontal rules.
"""

import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# -- Paths ----------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
DOCS_DIR = PROJECT_DIR / "docs"

# -- Colors (VCP palette) -------------------------------------
if _DOCX_AVAILABLE:
    NAVY = RGBColor(0x1B, 0x23, 0x33)
    TEAL = RGBColor(0x1A, 0xAF, 0xCC)
    GREEN = RGBColor(0x76, 0xB9, 0x00)
    GRAY_BODY = RGBColor(0x33, 0x33, 0x33)
    GRAY_META = RGBColor(0x66, 0x66, 0x66)
    GRAY_CODE = RGBColor(0x44, 0x44, 0x44)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_NAVY = "1B2333"
HEX_TEAL = "1AAFCC"
HEX_GREEN = "76B900"
HEX_LIGHT = "F8FAFB"
HEX_CODE_BG = "F0F0F0"
FONT = "Calibri"
CODE_FONT = "Courier New"

HEADER_TEXT = "HCLS AI Factory \u2014 Rare Disease Diagnostic Agent"
FOOTER_TEXT = "Confidential \u2014 For Internal Use"


# -- Helpers --------------------------------------------------

def _set_cell_shading(cell, hex_color: str) -> None:
    """Set background color on a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_run(paragraph, text: str, bold: bool = False, italic: bool = False,
             font_name: str = FONT, size: int = 11, color=None) -> None:
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _parse_inline(text: str):
    """Yield (text, bold, italic, is_code) tuples from Markdown inline formatting."""
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            yield (text[pos:m.start()], False, False, False)
        if m.group(2):  # bold italic
            yield (m.group(2), True, True, False)
        elif m.group(3):  # bold
            yield (m.group(3), True, False, False)
        elif m.group(4):  # italic
            yield (m.group(4), False, True, False)
        elif m.group(5):  # code
            yield (m.group(5), False, False, True)
        pos = m.end()
    if pos < len(text):
        yield (text[pos:], False, False, False)


# -- Main Converter -------------------------------------------

def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    """Convert a single Markdown file to a branded DOCX."""
    if not _DOCX_AVAILABLE:
        print(f"python-docx not installed -- skipping {md_path.name}")
        return

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Header
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(hp, HEADER_TEXT, size=8, color=GRAY_META)

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(fp, FOOTER_TEXT, size=8, color=GRAY_META, italic=True)

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    in_code = False
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            p = doc.add_paragraph()
            _add_run(p, line, font_name=CODE_FONT, size=9, color=GRAY_CODE)
            continue

        # Horizontal rules
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 60)
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            h = doc.add_heading(heading_text, level=min(level, 4))
            for run in h.runs:
                run.font.name = FONT
                run.font.color.rgb = NAVY if level <= 2 else TEAL
            continue

        # Tables
        if stripped.startswith("|"):
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                continue  # separator row
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_rows = [cells]
            else:
                table_rows.append(cells)
            continue
        elif in_table:
            # Flush table
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=ncols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row_data in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < ncols:
                            cell = tbl.cell(ri, ci)
                            cell.text = cell_text
                            if ri == 0:
                                _set_cell_shading(cell, HEX_NAVY)
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.font.color.rgb = WHITE
                                        run.font.bold = True
            in_table = False
            table_rows = []

        # Bullet lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            for txt, bold, italic, is_code in _parse_inline(content):
                if is_code:
                    _add_run(p, txt, font_name=CODE_FONT, size=10, color=GRAY_CODE)
                else:
                    _add_run(p, txt, bold=bold, italic=italic, color=GRAY_BODY)
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            content = num_match.group(2)
            p = doc.add_paragraph(style="List Number")
            for txt, bold, italic, is_code in _parse_inline(content):
                if is_code:
                    _add_run(p, txt, font_name=CODE_FONT, size=10, color=GRAY_CODE)
                else:
                    _add_run(p, txt, bold=bold, italic=italic, color=GRAY_BODY)
            continue

        # Blockquotes
        if stripped.startswith("> "):
            content = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            _add_run(p, content, italic=True, color=GRAY_META)
            continue

        # Empty lines
        if not stripped:
            doc.add_paragraph()
            continue

        # Normal paragraphs
        p = doc.add_paragraph()
        for txt, bold, italic, is_code in _parse_inline(stripped):
            if is_code:
                _add_run(p, txt, font_name=CODE_FONT, size=10, color=GRAY_CODE)
            else:
                _add_run(p, txt, bold=bold, italic=italic, color=GRAY_BODY)

    # Flush any remaining table
    if in_table and table_rows:
        ncols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=ncols)
        for ri, row_data in enumerate(table_rows):
            for ci, cell_text in enumerate(row_data):
                if ci < ncols:
                    tbl.cell(ri, ci).text = cell_text

    doc.save(str(docx_path))
    print(f"Generated: {docx_path}")


def main() -> None:
    """Convert all Markdown files in docs/ to DOCX."""
    if not DOCS_DIR.exists():
        print(f"No docs/ directory found at {DOCS_DIR}")
        return

    md_files = sorted(DOCS_DIR.glob("*.md"))
    if not md_files:
        print(f"No Markdown files found in {DOCS_DIR}")
        return

    success = 0
    errors = []

    for md_file in md_files:
        docx_path = DOCS_DIR / md_file.with_suffix(".docx").name
        try:
            markdown_to_docx(md_file, docx_path)
            size_kb = docx_path.stat().st_size / 1024
            print(f"  OK    {docx_path.name}  ({size_kb:.0f} KB)")
            success += 1
        except Exception as exc:
            msg = f"  FAIL  {md_file.name}: {exc}"
            print(msg)
            errors.append(msg)

    print(f"\nDone: {success}/{len(md_files)} files converted.")
    if errors:
        print("Errors:")
        for e in errors:
            print(f"  {e}")


if __name__ == "__main__":
    main()
